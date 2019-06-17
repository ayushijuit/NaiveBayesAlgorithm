import docx
import re
import inflect
import math
import sklearn
from nltk.tokenize import word_tokenize,sent_tokenize
from nltk.corpus import stopwords
from nltk.stem.porter import PorterStemmer
from nltk import pos_tag


contractions = { 
"ain't": "am not",
"aren't": "are not",
"can't": "cannot",
"can't've": "cannot have",
"'cause": "because",
"could've": "could have",
"couldn't": "could not",
"couldn't've": "could not have",
"didn't": "did not",
"doesn't": "does not",
"don't": "do not",
"hadn't": "had not",
"hadn't've": "had not have",
"hasn't": "has not",
"haven't": "have not",
"he'd": "he would",
"he'd've": "he would have",
"he'll": "he will",
"he's": "he is",
"how'd": "how did",
"how'll": "how will",
"how's": "how is",
"i'd": "i would",
"i'll": "i will",
"i'm": "i am",
"i've": "i have",
"isn't": "is not",
"it'd": "it would",
"it'll": "it will",
"it's": "it is",
"let's": "let us",
"ma'am": "madam",
"mayn't": "may not",
"might've": "might have",
"mightn't": "might not",
"must've": "must have",
"mustn't": "must not",
"needn't": "need not",
"oughtn't": "ought not",
"shan't": "shall not",
"sha'n't": "shall not",
"she'd": "she would",
"she'll": "she will",
"she's": "she is",
"should've": "should have",
"shouldn't": "should not",
"that'd": "that would",
"that's": "that is",
"there'd": "there had",
"there's": "there is",
"they'd": "they would",
"they'll": "they will",
"they're": "they are",
"they've": "they have",
"wasn't": "was not",
"we'd": "we would",
"we'll": "we will",
"we're": "we are",
"we've": "we have",
"weren't": "were not",
"what'll": "what will",
"what're": "what are",
"what's": "what is",
"what've": "what have",
"where'd": "where did",
"where's": "where is",
"who'll": "who will",
"who's": "who is",
"won't": "will not",
"wouldn't": "would not",
"you'd": "you would",
"you'll": "you will",
"you're": "you are"
}

def clean_text(text, remove_stopwords = True):
    '''Remove unwanted characters, and format the text to create fewer nulls word embeddings'''
    
    # Convert words to lower case
    text = text.lower()
    
    # Replace contractions with their longer forms 
    if True:
        # We are not using "text.split()" here
        #since it is not fool proof, e.g. words followed by punctuations "Are you kidding?I think you aren't."
        text = re.findall(r"[\w']+", text)
        new_text = []
        for word in text:
            if word in contractions:
                new_text.append(contractions[word])
            else:
                new_text.append(word)
        text = " ".join(new_text)
    
    # Format words and remove unwanted characters
    text = re.sub(r'https?:\/\/.*[\r\n]*', '', text, flags=re.MULTILINE)# remove links
    text = re.sub(r'\<a href', ' ', text)# remove html link tag
    text = re.sub(r'&amp;', '', text) 
    text = re.sub(r'[_"\-;%()|+&=*%.,!?:#$@\[\]/]', ' ', text)
    text = re.sub(r'<br />', ' ', text)
    text = re.sub(r'\'', ' ', text)
    
    return text


def tokenize(doc):
        wholedoc=""
        for para in doc.paragraphs:
                wholedoc=wholedoc+para.text
        wholedoc=clean_text(wholedoc, remove_stopwords=False)
        tokens=word_tokenize(wholedoc)
        return tokens

doc1= docx.Document('Doc1.docx')
doc2= docx.Document('Doc2.docx')
doc3= docx.Document('Doc3.docx')
doc4= docx.Document('Doc4.docx')
doc5= docx.Document('Doc5.docx')
doc6= docx.Document('Doc6.docx')
doc7= docx.Document('Doc7.docx')
doc8= docx.Document('Doc8.docx')
doc9= docx.Document('Doc9.docx')
doc10= docx.Document('Doc10.docx')
words1=tokenize(doc1)
words2=tokenize(doc2)
words3=tokenize(doc3)
words4=tokenize(doc4)
words5=tokenize(doc5)
words6=tokenize(doc6)
words7=tokenize(doc7)
words8=tokenize(doc8)
words9=tokenize(doc9)
words10=tokenize(doc10)
words=words1+words2+words3+words4+words5+words6+words7+words8+words9+words10

is_noun = lambda pos: pos[:2] == 'NN'
nouns = [word for (word, pos) in pos_tag(words) if is_noun(pos)]

class1= docx.Document('class1.docx')
class0= docx.Document('class0.docx')
wholedoc=""
for para in class1.paragraphs:
    wholedoc=wholedoc+para.text
class1_sentences=sent_tokenize(wholedoc)
class1_words=tokenize(class1)
len_class1=len(class1_sentences)
wholedoc=""
for para in class0.paragraphs:
    wholedoc=wholedoc+para.text
class0_sentences=sent_tokenize(wholedoc)
class0_words=tokenize(class0)
len_class0=len(class0_sentences)


word_frequencies_insummary={}
n1=0
for word in class1_words:
    if word in nouns:
        n1=n1+1
        if word not in word_frequencies_insummary.keys():
            word_frequencies_insummary[word] = 1
        else:
            word_frequencies_insummary[word] += 1
    else:
        word_frequencies_insummary[word] = 0

#print(word_frequencies_insummary)


word_frequencies_notinsummary={}

n2=0
for word in class0_words:
    if word in nouns:
        n2=n2+1
        if word not in word_frequencies_notinsummary.keys():
            word_frequencies_notinsummary[word] = 1
        else:
            word_frequencies_notinsummary[word] += 1
    else:
        word_frequencies_notinsummary[word] = 0
 
#print(word_frequencies_notinsummary)





test_doc= doc9
wholedoc=""
for para in test_doc.paragraphs:
    wholedoc=wholedoc+para.text
test_sentences = sent_tokenize(wholedoc)
print(test_sentences)
counts_class1={}
counts_class0={}
m=0
i=1
summary_sentences=[]
for sentence in test_sentences:
   
    w_class1=0
    w_class0=0
    for word in sentence.split():
        if word in nouns:
            if word in word_frequencies_insummary:
                counts_class1[word]=word_frequencies_insummary[word]
            else:
                counts_class1[word]=0
            if word in word_frequencies_notinsummary:
                counts_class0[word]=word_frequencies_notinsummary[word]
            else:
                counts_class0[word]=0
            m=m+1
            w_class1=w_class1+math.log((counts_class1[word]+1)/(n1+len(nouns)))
            w_class0=w_class0+math.log((counts_class0[word]+1)/(n2+len(nouns)))
    w_class1=w_class1+math.log((len_class1/(len_class1+len_class0)))+math.log(m/len_class1)+math.log((1/i)/len_class1)
    w_class0=w_class0+math.log((len_class0/(len_class1+len_class0)))+math.log(m/len_class0)+math.log((1/i)/len_class0)
    print("Sentence",i,"Score for Class1:",w_class1)
    print("Sentence",i,"Score for Class0:",w_class0)
    print("\n")
    i=i+1
    if(w_class1>w_class0):
        summary_sentences.append(sentence)
#print(len(summary_sentences))
summary = ' '.join(summary_sentences)
print(summary)

testdocclass1=docx.Document('testdocclass2.docx')
wholedoc=""
for para in testdocclass1.paragraphs:
    wholedoc=wholedoc+para.text
test_sentences = sent_tokenize(wholedoc)

tp=0
fp=0
fn=0
tn=0
for x in range(len(summary_sentences)):
    if (summary_sentences[x] in test_sentences):
        tp=tp+1
    else:
        fp=fp+1
pre=tp/(tp+fp)*100
print("Precision:",pre)

for x in range(len(test_sentences)):
    if(test_sentences[x] not in summary_sentences):
        fn=fn+1
rec=tp/(tp+fn)*100
print("Recall:",rec)

fscore=2*((pre*rec)/(pre+rec))
print("F-score:",fscore)
print(len(summary_sentences))
print(len(test_sentences))
print(tp)
print(fp)
print(fn)

